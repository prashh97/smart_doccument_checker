"""
Document Processing Module
Handles file upload, text extraction, and basic preprocessing
"""

import streamlit as st
from typing import List, Dict, Any, Optional
import io
from datetime import datetime
from pathlib import Path
import hashlib

# Document processing imports
try:
    import PyPDF2
    from docx import Document
    import pdfplumber
except ImportError as e:
    st.error(f"Missing dependency: {e}. Please install required packages.")

class DocumentProcessor:
    """Handles document upload and text extraction"""
    
    def __init__(self):
        self.supported_types = {
            "text/plain": self._extract_text_file,
            "application/pdf": self._extract_pdf_file,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._extract_docx_file,
            "application/msword": self._extract_doc_file,
            "text/markdown": self._extract_text_file,
            "text/html": self._extract_text_file,
        }
    
    def process_uploaded_files(self, uploaded_files: List[Any]) -> List[Dict[str, Any]]:
        """
        Process multiple uploaded files and extract text content
        
        Args:
            uploaded_files: List of Streamlit uploaded file objects
            
        Returns:
            List of document dictionaries with metadata
        """
        
        documents = []
        
        for uploaded_file in uploaded_files:
            try:
                doc_data = self._process_single_file(uploaded_file)
                if doc_data:
                    documents.append(doc_data)
                    
            except Exception as e:
                st.error(f"Error processing {uploaded_file.name}: {str(e)}")
                continue
        
        return documents
    
    def _process_single_file(self, uploaded_file: Any) -> Optional[Dict[str, Any]]:
        """Process a single uploaded file"""
        
        # Validate file
        if not self._validate_file(uploaded_file):
            return None
        
        # Extract text based on file type
        content = self._extract_text_content(uploaded_file)
        
        if not content:
            st.warning(f"Could not extract text from {uploaded_file.name}")
            return None
        
        # Create document metadata
        doc_data = {
            "name": uploaded_file.name,
            "content": content,
            "size": len(content),
            "file_size": uploaded_file.size,
            "file_type": uploaded_file.type,
            "word_count": len(content.split()),
            "char_count": len(content),
            "upload_time": datetime.now(),
            "file_hash": self._calculate_file_hash(content),
            "preview": content[:500] + "..." if len(content) > 500 else content
        }
        
        return doc_data
    
    def _validate_file(self, uploaded_file: Any) -> bool:
        """Validate uploaded file"""
        
        # Check file size (50MB limit)
        max_size = 50 * 1024 * 1024  # 50MB in bytes
        if uploaded_file.size > max_size:
            st.error(f"File {uploaded_file.name} is too large. Maximum size: 50MB")
            return False
        
        # Check file type
        if uploaded_file.type not in self.supported_types:
            st.error(f"Unsupported file type: {uploaded_file.type}")
            return False
        
        return True
    
    def _extract_text_content(self, uploaded_file: Any) -> str:
        """Extract text content based on file type"""
        
        file_type = uploaded_file.type
        extractor_func = self.supported_types.get(file_type)
        
        if not extractor_func:
            return ""
        
        try:
            # Reset file pointer
            uploaded_file.seek(0)
            content = extractor_func(uploaded_file)
            return content.strip() if content else ""
            
        except Exception as e:
            st.error(f"Error extracting text from {uploaded_file.name}: {str(e)}")
            return ""
    
    def _extract_text_file(self, uploaded_file: Any) -> str:
        """Extract text from plain text files"""
        try:
            content = uploaded_file.read().decode('utf-8', errors='ignore')
            return content
        except Exception:
            # Try different encodings
            for encoding in ['latin1', 'cp1252', 'ascii']:
                try:
                    uploaded_file.seek(0)
                    content = uploaded_file.read().decode(encoding, errors='ignore')
                    return content
                except Exception:
                    continue
            return ""
    
    def _extract_pdf_file(self, uploaded_file: Any) -> str:
        """Extract text from PDF files using multiple methods"""
        
        text_content = ""
        
        # Method 1: Try pdfplumber (better for complex PDFs)
        try:
            import pdfplumber
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
        except Exception:
            pass
        
        # Method 2: Fallback to PyPDF2 if pdfplumber fails
        if not text_content.strip():
            try:
                uploaded_file.seek(0)
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
            except Exception as e:
                st.error(f"PDF extraction failed: {str(e)}")
                return ""
        
        return text_content
    
    def _extract_docx_file(self, uploaded_file: Any) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(uploaded_file)
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            return text_content
            
        except Exception as e:
            st.error(f"DOCX extraction failed: {str(e)}")
            return ""
    
    def _extract_doc_file(self, uploaded_file: Any) -> str:
        """Extract text from DOC files (legacy format)"""
        # For .doc files, you might need python-docx2txt or other libraries
        # For now, return empty string with a message
        st.warning(f"Legacy .doc format not fully supported. Please convert {uploaded_file.name} to .docx")
        return ""
    
    def _calculate_file_hash(self, content: str) -> str:
        """Calculate hash of file content for deduplication"""
        return hashlib.md5(content.encode()).hexdigest()
    
    def get_document_stats(self, documents: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Calculate statistics for processed documents"""
        
        if not documents:
            return {}
        
        total_size = sum(doc['file_size'] for doc in documents)
        total_words = sum(doc['word_count'] for doc in documents)
        total_chars = sum(doc['char_count'] for doc in documents)
        
        file_types = {}
        for doc in documents:
            file_type = doc['file_type']
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        return {
            'total_documents': len(documents),
            'total_size_bytes': total_size,
            'total_size_mb': round(total_size / (1024 * 1024), 2),
            'total_words': total_words,
            'total_characters': total_chars,
            'average_words_per_doc': round(total_words / len(documents), 0),
            'file_types': file_types,
            'largest_document': max(documents, key=lambda x: x['word_count'])['name'],
            'smallest_document': min(documents, key=lambda x: x['word_count'])['name']
        }
    
    def detect_duplicate_documents(self, documents: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Detect duplicate documents based on content hash"""
        
        seen_hashes = {}
        duplicates = []
        
        for doc in documents:
            file_hash = doc['file_hash']
            if file_hash in seen_hashes:
                duplicates.append({
                    'original': seen_hashes[file_hash],
                    'duplicate': doc['name'],
                    'hash': file_hash
                })
            else:
                seen_hashes[file_hash] = doc['name']
        
        return duplicates
    
    def preprocess_text(self, text: str) -> str:
        """Basic text preprocessing for better analysis"""
        
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove non-printable characters
        text = ''.join(char for char in text if char.isprintable() or char.isspace())
        
        return text
    
    def extract_key_sections(self, content: str, document_name: str) -> Dict[str, str]:
        """Extract key sections from document content"""
        
        sections = {}
        
        # Common section headers to look for
        section_headers = [
            'policy', 'procedure', 'requirement', 'deadline', 'attendance',
            'submission', 'grade', 'grading', 'penalty', 'late', 'extension'
        ]
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
                
            # Check if paragraph contains key terms
            paragraph_lower = paragraph.lower()
            for header in section_headers:
                if header in paragraph_lower:
                    if header not in sections:
                        sections[header] = []
                    sections[header].append(paragraph[:200])  # First 200 chars
        
        return sections
    
    @staticmethod
    def get_current_timestamp() -> datetime:
        """Get current timestamp"""
        return datetime.now()
    
    def save_processed_documents(self, documents: List[Dict[str, Any]], analysis_id: str) -> str:
        """Save processed documents to cache for reuse"""
        
        import json
        from config.settings import AppSettings
        
        settings = AppSettings()
        cache_file = settings.CACHE_DIR / f"analysis_{analysis_id}.json"
        
        # Prepare data for saving (remove non-serializable items)
        save_data = []
        for doc in documents:
            doc_copy = doc.copy()
            doc_copy['upload_time'] = doc_copy['upload_time'].isoformat()
            save_data.append(doc_copy)
        
        try:
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(save_data, f, indent=2, ensure_ascii=False)
            return str(cache_file)
        except Exception as e:
            st.error(f"Failed to cache documents: {str(e)}")
            return ""
    
    def load_processed_documents(self, analysis_id: str) -> Optional[List[Dict[str, Any]]]:
        """Load processed documents from cache"""
        
        import json
        from config.settings import AppSettings
        
        settings = AppSettings()
        cache_file = settings.CACHE_DIR / f"analysis_{analysis_id}.json"
        
        try:
            if cache_file.exists():
                with open(cache_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # Convert timestamp strings back to datetime objects
                for doc in data:
                    doc['upload_time'] = datetime.fromisoformat(doc['upload_time'])
                
                return data
        except Exception as e:
            st.error(f"Failed to load cached documents: {str(e)}")
        
        return None